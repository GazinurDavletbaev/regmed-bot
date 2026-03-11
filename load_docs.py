import os
from pathlib import Path
from qdrant_client import QdrantClient
from qdrant_client.http import models
from sentence_transformers import SentenceTransformer
from pypdf import PdfReader
from docx import Document

# Конфигурация
QDRANT_HOST = "localhost"
QDRANT_PORT = 6333
COLLECTION_NAME = "med_docs"
DOCS_FOLDER = "./docs"  # папка с документами
CHUNK_SIZE = 500        # примерный размер чанка в символах
CHUNK_OVERLAP = 50      # перекрытие между чанками

# Подключение к Qdrant
client = QdrantClient(QDRANT_HOST, port=QDRANT_PORT)
model = SentenceTransformer('intfloat/multilingual-e5-large')

def extract_text_from_pdf(path):
    reader = PdfReader(path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(path):
    doc = Document(path)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    tables_text.append(cell_text)
    all_text = "\n".join(paragraphs + tables_text)
    print(f"Файл {path.name}: найдено абзацев: {len(paragraphs)}, ячеек таблиц: {len(tables_text)}")
    return all_text

def split_into_chunks(text, size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    words = text.split()
    chunks = []
    for i in range(0, len(words), size - overlap):
        chunk = " ".join(words[i:i+size])
        chunks.append(chunk)
    return chunks

# Собираем все файлы
files = list(Path(DOCS_FOLDER).glob("*.pdf")) + list(Path(DOCS_FOLDER).glob("*.docx"))
print(f"Найдено файлов: {len(files)}")

all_chunks = []
all_metadata = []

for file_path in files:
    print(f"Обработка: {file_path.name}")
    if file_path.suffix == ".pdf":
        text = extract_text_from_pdf(file_path)
        print(f"Длина текста файла {file_path.name}: {len(text)} символов")
    elif file_path.suffix == ".docx":
        text = extract_text_from_docx(file_path)
        print(f"Длина текста файла {file_path.name}: {len(text)} символов")
    else:
        continue

    chunks = split_into_chunks(text)
    for i, chunk in enumerate(chunks):
        all_chunks.append(chunk)
        all_metadata.append({
            "source": str(file_path.name),
            "chunk": i,
            "text": chunk
        })

print(f"Всего чанков: {len(all_chunks)}")

# Создаём эмбеддинги (можно делать батчами, но для простоты так)
print("Создание эмбеддингов...")
vectors = model.encode(all_chunks).tolist()

# Загружаем в Qdrant
points = []
for idx, (vector, meta) in enumerate(zip(vectors, all_metadata)):
    points.append(
        models.PointStruct(
            id=idx,  # для простоты, но лучше использовать уникальный ID (например, хеш)
            vector=vector,
            payload=meta
        )
    )

# Заливаем пачками по 100 штук (чтобы не перегружать)
BATCH_SIZE = 100
for i in range(0, len(points), BATCH_SIZE):
    client.upsert(
        collection_name=COLLECTION_NAME,
        points=points[i:i+BATCH_SIZE]
    )
    print(f"Загружено {i+BATCH_SIZE if i+BATCH_SIZE < len(points) else len(points)} из {len(points)}")

print("Готово!")